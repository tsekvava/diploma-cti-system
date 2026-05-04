"""Input loader for CTI reports."""

from __future__ import annotations

import base64
import io
import os
from urllib.parse import unquote, urlparse
from dataclasses import dataclass, field
from pathlib import Path


TEXT_EXTENSIONS = {".txt", ".md", ".log"}
HTML_EXTENSIONS = {".html", ".htm"}
DOCX_EXTENSIONS = {".docx"}
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".bmp"}


class InputLoadError(RuntimeError):
    """Ошибки парсинга входного файла отчёта."""


@dataclass
class LoadedReport:
    text: str
    source: str
    input_format: str
    pages: int = 1
    ocr_used: bool = False
    parse_warnings: list[str] = field(default_factory=list)


def load_report_file(file_path: str | Path, ocr_mode: str = "auto", ocr_lang: str = "rus+eng") -> LoadedReport:
    """Загружает отчёт из файла в plain text + метаданные парсинга."""
    path = Path(file_path)
    if not path.exists():
        raise InputLoadError(f"Файл не найден: {path}")

    ext = path.suffix.lower()
    warnings: list[str] = []
    source = path.name

    if ext in TEXT_EXTENSIONS:
        text = _read_text_with_fallback(path, warnings)
        pages = 1
        input_format = ext.lstrip(".")
        ocr_used = False
    elif ext in HTML_EXTENSIONS:
        raw = _read_text_with_fallback(path, warnings)
        text, html_ocr_used, html_warnings = _extract_html_text(
            raw,
            base_dir=path.parent,
            ocr_mode=ocr_mode,
            ocr_lang=ocr_lang,
        )
        warnings.extend(html_warnings)
        pages = 1
        input_format = "html"
        ocr_used = html_ocr_used
    elif ext in DOCX_EXTENSIONS:
        text, docx_ocr_used, docx_warnings = _extract_docx_text(
            path,
            ocr_mode=ocr_mode,
            ocr_lang=ocr_lang,
        )
        warnings.extend(docx_warnings)
        pages = 1
        input_format = "docx"
        ocr_used = docx_ocr_used
    elif ext in PDF_EXTENSIONS:
        text, pages, ocr_used, pdf_warnings = _extract_pdf_text(path, ocr_mode=ocr_mode, ocr_lang=ocr_lang)
        warnings.extend(pdf_warnings)
        input_format = "pdf"
    elif ext in IMAGE_EXTENSIONS:
        text, img_warnings = _extract_image_text(path, ocr_lang=ocr_lang)
        warnings.extend(img_warnings)
        pages = 1
        input_format = "image"
        ocr_used = True
    else:
        supported = ", ".join(sorted(TEXT_EXTENSIONS | HTML_EXTENSIONS | DOCX_EXTENSIONS | PDF_EXTENSIONS | IMAGE_EXTENSIONS))
        raise InputLoadError(
            f"Неподдерживаемый формат: {ext or '[no extension]'}.\n"
            f"Поддерживаемые: {supported}"
        )

    text = (text or "").strip()
    if not text:
        details = f" Warnings: {'; '.join(warnings)}" if warnings else ""
        raise InputLoadError(
            f"Не удалось извлечь текст из файла '{source}'. "
            f"Проверь формат, кодировку и наличие OCR (если это скан/PDF-изображение)."
            f"{details}"
        )

    return LoadedReport(
        text=text,
        source=source,
        input_format=input_format,
        pages=pages,
        ocr_used=ocr_used,
        parse_warnings=warnings,
    )


def _read_text_with_fallback(path: Path, warnings: list[str]) -> str:
    raw = path.read_bytes()
    for enc in ("utf-8", "cp1251"):
        try:
            text = raw.decode(enc)
            if enc != "utf-8":
                warnings.append(f"Файл '{path.name}' прочитан с fallback-кодировкой {enc}.")
            return text
        except UnicodeDecodeError:
            continue

    warnings.append(f"Файл '{path.name}' содержит нестандартную кодировку. Применён decode(errors='replace').")
    return raw.decode("utf-8", errors="replace")


def _extract_html_text(
    html: str,
    base_dir: Path,
    ocr_mode: str,
    ocr_lang: str,
) -> tuple[str, bool, list[str]]:
    warnings: list[str] = []
    text = ""
    image_sources: list[str] = []

    try:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        text = soup.get_text(separator="\n", strip=True)
        image_sources = [
            str(img.get("src", "")).strip()
            for img in soup.find_all("img")
            if str(img.get("src", "")).strip()
        ]
    except Exception:
        # Fallback на стандартный HTMLParser, если bs4 не установлен.
        from html.parser import HTMLParser

        class _TextExtractor(HTMLParser):
            def __init__(self):
                super().__init__()
                self.parts = []
                self.images = []

            def handle_data(self, data):
                value = str(data or "").strip()
                if value:
                    self.parts.append(value)

            def handle_starttag(self, tag, attrs):
                if str(tag).lower() != "img":
                    return
                attrs_map = dict(attrs or [])
                src = str(attrs_map.get("src", "")).strip()
                if src:
                    self.images.append(src)

        parser = _TextExtractor()
        parser.feed(html)
        text = "\n".join(parser.parts)
        image_sources = parser.images

    ocr_chunks: list[str] = []
    ocr_used = False
    if image_sources and _should_use_ocr(text, ocr_mode):
        ocr_chunks = _ocr_html_images(
            image_sources=image_sources,
            base_dir=base_dir,
            ocr_lang=ocr_lang,
            warnings=warnings,
        )
        if ocr_chunks:
            ocr_used = True

    if ocr_chunks:
        text = _merge_text_and_ocr(text, ocr_chunks)

    return text, ocr_used, warnings


def _extract_docx_text(path: Path, ocr_mode: str, ocr_lang: str) -> tuple[str, bool, list[str]]:
    warnings: list[str] = []

    try:
        from docx import Document
    except Exception as e:
        raise InputLoadError(
            "Для .docx нужен python-docx. "
            "Установи: pip install python-docx"
        ) from e

    doc = Document(str(path))
    chunks: list[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [(cell.text or "").strip() for cell in row.cells]
            line = " | ".join([c for c in cells if c])
            if line:
                chunks.append(line)

    text = "\n".join(chunks)
    ocr_chunks: list[str] = []
    ocr_used = False
    if _should_use_ocr(text, ocr_mode):
        ocr_chunks = _ocr_docx_images(doc, ocr_lang=ocr_lang, warnings=warnings)
        if ocr_chunks:
            ocr_used = True

    if ocr_chunks:
        text = _merge_text_and_ocr(text, ocr_chunks)

    return text, ocr_used, warnings


def _extract_pdf_text(path: Path, ocr_mode: str, ocr_lang: str) -> tuple[str, int, bool, list[str]]:
    warnings: list[str] = []

    try:
        from pypdf import PdfReader
    except Exception as e:
        raise InputLoadError("Для .pdf нужен pypdf. Установи: pip install pypdf") from e

    try:
        reader = PdfReader(str(path))
    except Exception as e:
        raise InputLoadError(f"Не удалось открыть PDF '{path.name}': {e}") from e

    total_pages = len(reader.pages)
    page_texts: list[str] = []
    ocr_used = False
    pdfium_doc = None

    for idx, page in enumerate(reader.pages):
        extracted = ""
        try:
            extracted = page.extract_text() or ""
        except Exception as e:
            warnings.append(f"PDF page {idx + 1}: ошибка text extraction ({e}).")

        extracted = _normalize_text(extracted)
        need_ocr = _should_use_ocr(extracted, ocr_mode)

        if need_ocr:
            ocr_text = _ocr_pdf_page(path, idx, ocr_lang, warnings, pdfium_doc)
            if ocr_text is not None:
                if pdfium_doc is None:
                    pdfium_doc = ocr_text[1]
                ocr_value = _normalize_text(ocr_text[0])
                if ocr_value:
                    ocr_used = True
                    if ocr_mode == "on" or len(ocr_value) > len(extracted):
                        extracted = ocr_value

        if extracted:
            page_texts.append(extracted)

    if pdfium_doc is not None and hasattr(pdfium_doc, "close"):
        try:
            pdfium_doc.close()
        except Exception:
            pass

    text = "\n\n".join(page_texts).strip()
    if not text:
        warnings.append("PDF не дал извлекаемого текста. Проверь OCR/Tesseract.")

    return text, total_pages, ocr_used, warnings


def _should_use_ocr(text: str, ocr_mode: str) -> bool:
    mode = (ocr_mode or "auto").lower()
    if mode == "off":
        return False
    if mode == "on":
        return True

    # auto
    cleaned = text.strip()
    if len(cleaned) < 80:
        return True

    alnum = sum(1 for ch in cleaned if ch.isalnum())
    ratio = alnum / max(1, len(cleaned))
    return ratio < 0.35


def _ocr_pdf_page(path: Path, page_idx: int, ocr_lang: str, warnings: list[str], pdfium_doc):
    try:
        import pypdfium2 as pdfium
    except Exception:
        warnings.append("OCR для PDF недоступен: нет pypdfium2 (pip install pypdfium2).")
        return None

    if pdfium_doc is None:
        try:
            pdfium_doc = pdfium.PdfDocument(str(path))
        except Exception as e:
            warnings.append(f"OCR для PDF недоступен: не удалось открыть через pypdfium2 ({e}).")
            return None

    try:
        page = pdfium_doc[page_idx]
        render = page.render(scale=2.0)
        image = render.to_pil()
        text = _ocr_image_obj(image, ocr_lang, warnings)
        if hasattr(page, "close"):
            page.close()
        return text, pdfium_doc
    except Exception as e:
        warnings.append(f"PDF page {page_idx + 1}: OCR error ({e}).")
        return None


def _extract_image_text(path: Path, ocr_lang: str) -> tuple[str, list[str]]:
    warnings: list[str] = []
    try:
        from PIL import Image
    except Exception as e:
        raise InputLoadError("Для OCR изображений нужен Pillow. Установи: pip install pillow") from e

    try:
        image = Image.open(str(path))
    except Exception as e:
        raise InputLoadError(f"Не удалось открыть изображение '{path.name}': {e}") from e

    text = _ocr_image_obj(image, ocr_lang, warnings)
    return _normalize_text(text), warnings


def _ocr_docx_images(doc, ocr_lang: str, warnings: list[str]) -> list[str]:
    blobs: list[bytes] = []
    seen_hashes: set[int] = set()

    try:
        rels = getattr(doc.part, "rels", {})
        for rel in rels.values():
            reltype = str(getattr(rel, "reltype", ""))
            if "/image" not in reltype:
                continue
            target_part = getattr(rel, "target_part", None)
            blob = getattr(target_part, "blob", None)
            if not blob:
                continue
            marker = hash(blob)
            if marker in seen_hashes:
                continue
            seen_hashes.add(marker)
            blobs.append(blob)
    except Exception as e:
        warnings.append(f"DOCX OCR warning: не удалось извлечь изображения ({e}).")
        return []

    ocr_chunks: list[str] = []
    for idx, blob in enumerate(blobs, start=1):
        text = _ocr_image_bytes(blob, ocr_lang=ocr_lang, warnings=warnings, source=f"docx-image-{idx}")
        normalized = _normalize_text(text)
        if normalized:
            ocr_chunks.append(normalized)

    return ocr_chunks


def _ocr_html_images(image_sources: list[str], base_dir: Path, ocr_lang: str, warnings: list[str]) -> list[str]:
    ocr_chunks: list[str] = []

    for idx, src in enumerate(image_sources, start=1):
        image_bytes = _load_html_image_bytes(src=src, base_dir=base_dir, warnings=warnings)
        if not image_bytes:
            continue
        text = _ocr_image_bytes(image_bytes, ocr_lang=ocr_lang, warnings=warnings, source=f"html-image-{idx}")
        normalized = _normalize_text(text)
        if normalized:
            ocr_chunks.append(normalized)

    return ocr_chunks


def _load_html_image_bytes(src: str, base_dir: Path, warnings: list[str]) -> bytes | None:
    value = (src or "").strip()
    if not value:
        return None

    if value.startswith("data:image"):
        marker = "base64,"
        if marker not in value:
            warnings.append("HTML OCR warning: data:image без base64 не поддерживается.")
            return None
        payload = value.split(marker, 1)[1].strip()
        if not payload:
            return None
        try:
            return base64.b64decode(payload, validate=False)
        except Exception as e:
            warnings.append(f"HTML OCR warning: не удалось декодировать data:image ({e}).")
            return None

    parsed = urlparse(value)
    scheme = parsed.scheme.lower()
    if scheme in ("http", "https"):
        warnings.append("HTML OCR warning: удалённые изображения по http/https не загружаются.")
        return None

    if scheme == "file":
        image_path = Path(unquote(parsed.path))
    elif scheme:
        warnings.append(f"HTML OCR warning: неподдерживаемая схема изображения '{scheme}'.")
        return None
    else:
        clean_src = value.split("#", 1)[0].split("?", 1)[0]
        rel_path = Path(unquote(clean_src))
        image_path = rel_path if rel_path.is_absolute() else (base_dir / rel_path)

    if not image_path.exists():
        warnings.append(f"HTML OCR warning: файл изображения не найден ({image_path}).")
        return None

    try:
        return image_path.read_bytes()
    except Exception as e:
        warnings.append(f"HTML OCR warning: не удалось прочитать изображение '{image_path.name}' ({e}).")
        return None


def _ocr_image_bytes(image_bytes: bytes, ocr_lang: str, warnings: list[str], source: str) -> str:
    try:
        from PIL import Image
    except Exception:
        warnings.append("OCR недоступен: нет Pillow (pip install pillow).")
        return ""

    try:
        with Image.open(io.BytesIO(image_bytes)) as image:
            return _ocr_image_obj(image, ocr_lang, warnings)
    except Exception as e:
        warnings.append(f"OCR warning: не удалось открыть {source} ({e}).")
        return ""


def _ocr_image_obj(image, ocr_lang: str, warnings: list[str]) -> str:
    try:
        import pytesseract
        from pytesseract import TesseractNotFoundError
    except Exception:
        warnings.append("OCR недоступен: нет pytesseract (pip install pytesseract).")
        return ""

    tesseract_cmd = os.getenv("TESSERACT_CMD")
    if tesseract_cmd:
        pytesseract.pytesseract.tesseract_cmd = tesseract_cmd

    try:
        _ = pytesseract.get_tesseract_version()
    except TesseractNotFoundError:
        warnings.append(
            "OCR недоступен: не найден бинарник tesseract. "
            "Установи Tesseract и/или задай TESSERACT_CMD."
        )
        return ""
    except Exception as e:
        warnings.append(f"OCR warning: не удалось проверить tesseract ({e}).")

    try:
        return pytesseract.image_to_string(image, lang=ocr_lang or "rus+eng")
    except Exception as e:
        warnings.append(f"OCR error: {e}")
        return ""


def _normalize_text(text: str) -> str:
    return "\n".join([line.strip() for line in str(text or "").splitlines() if line.strip()]).strip()


def _merge_text_and_ocr(text: str, ocr_chunks: list[str]) -> str:
    base = _normalize_text(text)
    ocr_part = _normalize_text("\n".join(ocr_chunks))

    if base and ocr_part:
        return f"{base}\n\n{ocr_part}"
    if base:
        return base
    return ocr_part
